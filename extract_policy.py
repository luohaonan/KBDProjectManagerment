from __future__ import annotations

from pathlib import Path


def main() -> None:
    src = Path("policy.docx")
    if not src.exists():
        raise SystemExit(f"Missing {src.resolve()}")

    try:
        import docx  # type: ignore
    except Exception as e:  # pragma: no cover
        raise SystemExit(
            "python-docx not installed. Install with: pip install python-docx"
        ) from e

    d = docx.Document(str(src))

    out_lines: list[str] = []
    out_lines.append(f"# Extracted from: {src.name}")

    # Paragraphs
    out_lines.append("\n## Paragraphs\n")
    for p in d.paragraphs:
        t = (p.text or "").strip()
        if t:
            out_lines.append(t)

    # Tables (flattened)
    if d.tables:
        out_lines.append("\n## Tables\n")
    for ti, table in enumerate(d.tables, start=1):
        out_lines.append(f"\n### Table {ti}\n")
        for row in table.rows:
            cells = [" ".join((c.text or "").split()) for c in row.cells]
            # Skip empty rows
            if any(cells):
                out_lines.append(" | ".join(cells))

    Path("policy_extracted.md").write_text("\n".join(out_lines) + "\n", encoding="utf-8")
    print("Wrote policy_extracted.md")


if __name__ == "__main__":
    main()

